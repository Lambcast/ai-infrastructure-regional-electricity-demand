"""
build_table1_docx.py
Produces a paper-ready Table 1 (Descriptive Statistics) as a .docx file
formatted to NBER working paper conventions.

Output: outputs/tables/table1_descriptive_stats.docx

Run from project root:
    python scripts/build_table1_docx.py
"""

import pandas as pd
import numpy as np
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Step 1: Load data ────────────────────────────────────────────────────────
print("Step 1: Loading data")
df_controls = pd.read_csv("data/panel_with_controls.csv")
df_shape    = pd.read_csv("data/panel_load_shape.csv")
print(f"  panel_with_controls: {len(df_controls)} rows")
print(f"  panel_load_shape:    {len(df_shape)} rows")

# Merge min demand onto controls panel
df = df_controls.merge(
    df_shape[["ba", "year_month", "demand_min"]],
    on=["ba", "year_month"],
    how="left"
)
print(f"  After merge: {len(df)} rows, demand_min missing: {df['demand_min'].isna().sum()}")

# Rescale GDP from millions to billions for readable table display
df["gdp_billions"] = df["gdp"] / 1000
print(f"  GDP rescaled to billions: mean = {df['gdp_billions'].mean():,.1f}")

# ─── Step 2: Define variable order and labels ─────────────────────────────────
print("\nStep 2: Checking variable availability")

VARS = [
    ("avg_demand_mwh",       "Avg. hourly demand (MWh)"),
    ("demand_min",           "Min. hourly demand (MWh)"),
    ("queue_mw_filed_large", "Queue MW filed, \u2265100 MW"),
    ("queue_mw_active",      "Queue MW active (stock)"),
    ("hdd",                  "Heating degree days"),
    ("cdd",                  "Cooling degree days"),
    ("gdp_billions",         "Regional GDP (billions $)"),
]

VARS = [(c, l) for c, l in VARS if c in df.columns]
for col, label in VARS:
    print(f"  ✓ {col}")

# ─── Step 3: Compute stats ────────────────────────────────────────────────────
print("\nStep 3: Computing statistics")

def stats_block(data, var_list):
    rows = []
    for col, label in var_list:
        s = data[col].dropna()
        rows.append({
            "label": label,
            "N":     int(len(s)),
            "Mean":  s.mean(),
            "SD":    s.std(),
            "Min":   s.min(),
            "Max":   s.max(),
        })
    return rows

BA_LABELS = {
    "ERCO": "Panel A: ERCOT",
    "PJM":  "Panel B: PJM",
    "MISO": "Panel C: MISO",
}

panels = []
for ba_code, panel_title in BA_LABELS.items():
    sub = df[df["ba"] == ba_code]
    if len(sub) == 0:
        print(f"  WARNING: no rows for ba={ba_code}")
        continue
    panels.append((panel_title, stats_block(sub, VARS)))
    print(f"  {ba_code}: {len(sub)} rows")

panels.append(("Panel D: All Balancing Authorities", stats_block(df, VARS)))
print(f"  All BAs: {len(df)} rows")

# ─── Step 4: Formatting helpers ───────────────────────────────────────────────

def fmt_num(val, decimals=1):
    if pd.isna(val):
        return "—"
    return f"{val:,.{decimals}f}"

def set_cell_border(cell, sides):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def clear_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "FFFFFF")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def apply_row_rule(row, top=False, bottom=False):
    for cell in row.cells:
        clear_cell_borders(cell)
        sides = []
        if top:    sides.append("top")
        if bottom: sides.append("bottom")
        if sides:
            set_cell_border(cell, sides)

def write_cell(cell, text, bold=False, italic=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(text)
    run.bold       = bold
    run.italic     = italic
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)

# ─── Step 5: Build document ───────────────────────────────────────────────────
print("\nStep 4: Building document")
doc = Document()

section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# Table title
title_p   = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_p.paragraph_format.space_after = Pt(4)
title_run = title_p.add_run("Table 1. Descriptive Statistics")
title_run.bold = True
title_run.font.name = "Times New Roman"
title_run.font.size = Pt(11)

# Column definitions
COLS       = ["Variable", "N", "Mean", "Std. Dev.", "Min", "Max"]
COL_WIDTHS = [Inches(2.9), Inches(0.4), Inches(0.85),
              Inches(0.85), Inches(0.85), Inches(0.75)]

tbl = doc.add_table(rows=0, cols=len(COLS))
tbl.style = "Table Grid"

# ── Header row ────────────────────────────────────────────────────────────────
hdr = tbl.add_row()
for i, (heading, width) in enumerate(zip(COLS, COL_WIDTHS)):
    cell  = hdr.cells[i]
    cell.width = width
    align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
    write_cell(cell, heading, bold=True, align=align)
apply_row_rule(hdr, top=True, bottom=True)

# ── Panels ────────────────────────────────────────────────────────────────────
for p_idx, (panel_title, rows) in enumerate(panels):

    # Panel header — merged across all columns
    ph = tbl.add_row()
    merged = ph.cells[0]
    for j in range(1, len(COLS)):
        merged = merged.merge(ph.cells[j])
    pf = merged.paragraphs[0]
    pf.paragraph_format.space_before = Pt(5) if p_idx > 0 else Pt(2)
    pf.paragraph_format.space_after  = Pt(1)
    pr = pf.add_run(panel_title)
    pr.bold = True
    pr.italic = False
    pr.font.name = "Times New Roman"
    pr.font.size = Pt(10)
    apply_row_rule(ph, top=False, bottom=False)
    clear_cell_borders(merged)

    # Data rows
    for r in rows:
        dr = tbl.add_row()
        dr.height = Pt(14)

        # Variable label — indented with spaces
        write_cell(dr.cells[0], "   " + r["label"],
                   align=WD_ALIGN_PARAGRAPH.LEFT)

        # Stats — pick decimal places per variable
        label_lower = r["label"].lower()
        if "gdp" in label_lower:
            decimals = 1   # billions: e.g. 2,007.4
        elif "degree days" in label_lower:
            decimals = 0   # whole numbers fine for HDD/CDD
        elif "mwh" in label_lower or "mw" in label_lower:
            decimals = 0   # demand and queue in whole MW
        else:
            decimals = 1

        for col_idx, key in enumerate(["N", "Mean", "SD", "Min", "Max"], start=1):
            raw = r[key]
            if key == "N":
                txt = f"{int(raw):,}"
            else:
                txt = fmt_num(raw, decimals)
            write_cell(dr.cells[col_idx], txt,
                       align=WD_ALIGN_PARAGRAPH.RIGHT)

        apply_row_rule(dr, top=False, bottom=False)

# Bottom rule on final row
apply_row_rule(tbl.rows[-1], top=False, bottom=True)

# ── Notes ─────────────────────────────────────────────────────────────────────
notes_p = doc.add_paragraph()
notes_p.paragraph_format.space_before = Pt(6)
nr = notes_p.add_run(
    "Notes: Sample period January 2019\u2013December 2025. Unit of observation is "
    "balancing authority \u00d7 calendar month (84 months per BA; 252 total observations). "
    "Avg. hourly demand is the mean of all hourly EIA Form 930 readings within a BA-month. "
    "Min. hourly demand is the lowest single hourly reading within a BA-month. "
    "Queue MW filed (\u2265100 MW) is total megawatts of new interconnection requests "
    "filed by projects at or above 100 MW in a given BA-month, from LBNL Queued Up (2024 edition). "
    "Queue MW active (stock) is the cumulative megawatts ever filed through a given month. "
    "HDD and CDD from NOAA NCEI. "
    "Regional GDP is annual real GDP by state from BEA regional accounts (2017 dollars), "
    "assigned to balancing authorities by geographic crosswalk. "
    "MISO queue filed standard deviation reflects a large filing spike in 2022 associated "
    "with queue reform dynamics rather than underlying demand pressure; see Section 4."
)
nr.italic     = True
nr.font.name  = "Times New Roman"
nr.font.size  = Pt(9)

# ─── Step 6: Save ────────────────────────────────────────────────────────────
print("\nStep 5: Saving")
os.makedirs("outputs/tables", exist_ok=True)
out_path = "outputs/tables/table1_descriptive_stats.docx"
doc.save(out_path)
print(f"  Saved to {out_path}")
print("\nDone.")